import os
import sys
import json
import logging
from pathlib import Path
from typing import Dict, Any, Optional, List, Tuple
from datetime import datetime
import shutil
import base64
import re
from PIL import Image
import io
import tempfile
import subprocess

# Document processing imports
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import mammoth
import fitz  # PyMuPDF
from pdf2image import convert_from_path
import pytesseract
from bs4 import BeautifulSoup
import cssutils

# Windows-specific imports
try:
    import win32com.client
    import pythoncom

    WINDOWS_AVAILABLE = True
except ImportError:
    WINDOWS_AVAILABLE = False

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Suppress cssutils warnings
cssutils.log.setLevel(logging.ERROR)


class HighQualityHTMLConverter:
    """Convert documents to high-quality HTML with preserved formatting"""

    def __init__(self,
                 output_dir: str = "html_output",
                 enable_ocr: bool = True,
                 image_quality: int = 85,
                 pdf_dpi: int = 150):
        """
        Initialize converter with configuration options

        Args:
            output_dir: Directory for HTML output
            enable_ocr: Enable OCR for scanned PDFs
            image_quality: JPEG quality for embedded images (1-100)
            pdf_dpi: DPI for PDF to image conversion
        """
        self.output_dir = output_dir
        self.enable_ocr = enable_ocr
        self.image_quality = image_quality
        self.pdf_dpi = pdf_dpi

        # Create output directories
        os.makedirs(output_dir, exist_ok=True)
        self.assets_dir = os.path.join(output_dir, "assets")
        os.makedirs(self.assets_dir, exist_ok=True)

        # HTML template with responsive design
        self.html_template = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        {styles}
    </style>
</head>
<body>
    <div class="resume-container">
        {content}
    </div>
    <script>
        {scripts}
    </script>
</body>
</html>"""

        # Base CSS for all resumes
        self.base_css = """
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }

        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            background-color: #f5f5f5;
            padding: 20px;
        }

        .resume-container {
            max-width: 850px;
            margin: 0 auto;
            background-color: white;
            box-shadow: 0 0 20px rgba(0,0,0,0.1);
            padding: 40px;
            position: relative;
        }

        @media screen and (max-width: 768px) {
            body { padding: 10px; }
            .resume-container { padding: 20px; }
        }

        @media print {
            body { background-color: white; padding: 0; }
            .resume-container { box-shadow: none; max-width: 100%; padding: 0; }
        }

        /* Typography */
        h1 { font-size: 2.5em; margin-bottom: 0.5em; color: #2c3e50; }
        h2 { font-size: 1.8em; margin-top: 1em; margin-bottom: 0.5em; color: #34495e; border-bottom: 2px solid #3498db; padding-bottom: 0.3em; }
        h3 { font-size: 1.4em; margin-top: 0.8em; margin-bottom: 0.4em; color: #34495e; }
        h4 { font-size: 1.2em; margin-top: 0.6em; margin-bottom: 0.3em; color: #34495e; }
        h5, h6 { font-size: 1.1em; margin-top: 0.5em; margin-bottom: 0.3em; color: #34495e; }

        p { margin-bottom: 0.8em; text-align: justify; }

        /* Lists */
        ul, ol { margin-left: 20px; margin-bottom: 0.8em; }
        li { margin-bottom: 0.3em; }

        /* Tables */
        table { width: 100%; border-collapse: collapse; margin-bottom: 1em; }
        th, td { padding: 8px 12px; text-align: left; border-bottom: 1px solid #ddd; }
        th { background-color: #f8f9fa; font-weight: bold; }

        /* Images */
        img { max-width: 100%; height: auto; display: block; margin: 1em auto; }
        .inline-image { display: inline-block; vertical-align: middle; margin: 0 5px; }

        /* Text formatting */
        strong, b { font-weight: 600; color: #2c3e50; }
        em, i { font-style: italic; }
        u { text-decoration: underline; }

        /* Links */
        a { color: #3498db; text-decoration: none; }
        a:hover { text-decoration: underline; }

        /* Special sections */
        .header-section { text-align: center; margin-bottom: 2em; }
        .contact-info { margin-bottom: 1em; }
        .section { margin-bottom: 2em; }

        /* Preserve original spacing */
        .preserve-space { white-space: pre-wrap; }

        /* Page breaks for printing */
        .page-break { page-break-after: always; }

        /* Highlighted text */
        .highlight { background-color: #fffacd; padding: 2px 4px; }

        /* Custom fonts embedded */
        @font-face {
            font-family: 'ResumeFont';
            src: local('Arial'), local('Helvetica');
        }
        """

        # JavaScript for interactive features
        self.base_js = """
        // Enable smooth scrolling
        document.querySelectorAll('a[href^="#"]').forEach(anchor => {
            anchor.addEventListener('click', function (e) {
                e.preventDefault();
                document.querySelector(this.getAttribute('href')).scrollIntoView({
                    behavior: 'smooth'
                });
            });
        });

        // Print functionality
        function printResume() {
            window.print();
        }

        // Download as PDF (requires backend support or browser print-to-PDF)
        function downloadAsPDF() {
            window.print();
        }
        """

    def convert_docx_to_html(self, file_path: str) -> str:
        """Convert DOCX to HTML preserving formatting"""
        try:
            # Use mammoth for HTML conversion with custom style mappings
            style_map = """
                p[style-name='Heading 1'] => h1:fresh
                p[style-name='Heading 2'] => h2:fresh
                p[style-name='Heading 3'] => h3:fresh
                p[style-name='Heading 4'] => h4:fresh
                p[style-name='Heading 5'] => h5:fresh
                p[style-name='Heading 6'] => h6:fresh
                p[style-name='Title'] => h1:fresh
                p[style-name='Subtitle'] => h2:fresh
                r[style-name='Strong'] => strong
                r[style-name='Emphasis'] => em
                p[style-name='List Paragraph'] => li:fresh
                p[style-name='Quote'] => blockquote:fresh
            """

            with open(file_path, "rb") as docx_file:
                result = mammoth.convert_to_html(
                    docx_file,
                    style_map=style_map
                )

            html_content = result.value

            # Check for any conversion messages/warnings
            if result.messages:
                for message in result.messages:
                    logger.debug(f"Mammoth message: {message}")

            # If mammoth produced empty or minimal content, try fallback
            if not html_content or len(html_content.strip()) < 50:
                logger.info("Mammoth produced minimal output, using fallback")
                return self._fallback_docx_conversion(file_path)

            # Enhance with python-docx for additional details
            try:
                doc = Document(file_path)
                enhanced_html = self._enhance_docx_html(doc, html_content)
                return enhanced_html
            except Exception as e:
                logger.warning(f"Could not enhance HTML: {e}")
                return html_content

        except Exception as e:
            logger.error(f"Error converting DOCX: {e}")
            # Fallback to basic conversion
            return self._fallback_docx_conversion(file_path)

    def _enhance_docx_html(self, doc: Document, base_html: str) -> str:
        """Enhance HTML with additional formatting from python-docx"""
        soup = BeautifulSoup(base_html, 'html.parser')

        # Extract and embed images
        image_counter = 0
        for rel in doc.part.rels.values():
            if "image" in str(rel.target_ref):
                try:
                    image_data = rel.target_part.blob
                    image_ext = str(rel.target_ref).split('.')[-1].lower()
                    if image_ext in ['jpeg', 'jpg']:
                        image_ext = 'jpeg'
                    elif image_ext == 'png':
                        image_ext = 'png'
                    else:
                        continue  # Skip unsupported image formats

                    image_b64 = base64.b64encode(image_data).decode()

                    # Create img tag
                    img_tag = soup.new_tag('img')
                    img_tag['src'] = f"data:image/{image_ext};base64,{image_b64}"
                    img_tag['alt'] = f"Image {image_counter}"
                    img_tag['class'] = "embedded-image"

                    # Try to find appropriate place to insert
                    # This is simplified - in production you'd match image positions
                    body = soup.find('body') or soup
                    body.append(img_tag)

                    image_counter += 1
                except Exception as e:
                    logger.warning(f"Failed to embed image: {e}")

        # Extract tables with formatting if not already in HTML
        if not soup.find('table') and doc.tables:
            for table in doc.tables:
                table_html = self._convert_table_to_html(table)
                soup.append(BeautifulSoup(table_html, 'html.parser'))

        return str(soup)

    def _convert_table_to_html(self, table) -> str:
        """Convert Word table to HTML"""
        html = '<table class="docx-table">\n'

        for row in table.rows:
            html += '<tr>\n'
            for cell in row.cells:
                cell_text = cell.text.strip()
                # Check if it's a header row (simplified check)
                tag = 'th' if row == table.rows[0] else 'td'
                html += f'<{tag}>{cell_text}</{tag}>\n'
            html += '</tr>\n'

        html += '</table>\n'
        return html

    def _fallback_docx_conversion(self, file_path: str) -> str:
        """Fallback conversion using python-docx only"""
        try:
            doc = Document(file_path)
            html_parts = []

            # Process paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    # Safely get style name
                    style_name = 'Normal'
                    if para.style and hasattr(para.style, 'name'):
                        style_name = para.style.name

                    # Determine paragraph type based on style
                    if 'Heading' in style_name and len(style_name) > 7 and style_name[-1].isdigit():
                        level = style_name[-1]
                        html_parts.append(f'<h{level}>{self._escape_html(para.text)}</h{level}>')
                    elif style_name == 'Title':
                        html_parts.append(f'<h1>{self._escape_html(para.text)}</h1>')
                    elif style_name == 'Subtitle':
                        html_parts.append(f'<h2>{self._escape_html(para.text)}</h2>')
                    else:
                        # Check alignment
                        align_class = ''
                        try:
                            if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                                align_class = ' style="text-align: center;"'
                            elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                                align_class = ' style="text-align: right;"'
                            elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                                align_class = ' style="text-align: justify;"'
                        except:
                            # If alignment check fails, just continue without alignment
                            pass

                        # Process runs safely
                        try:
                            processed_text = self._process_runs(para.runs)
                        except:
                            # If run processing fails, just use plain text
                            processed_text = self._escape_html(para.text)

                        html_parts.append(f'<p{align_class}>{processed_text}</p>')

            # Process tables
            try:
                for table in doc.tables:
                    html_parts.append(self._convert_table_to_html(table))
            except Exception as e:
                logger.warning(f"Could not process tables: {e}")

            # If we got some content, return it
            if html_parts:
                return '\n'.join(html_parts)
            else:
                # If no content extracted, return a simple message
                return "<p>Document appears to be empty or could not be processed.</p>"

        except Exception as e:
            logger.error(f"Fallback DOCX conversion failed: {e}")
            return f"<p>Error converting document: {e}</p>"

    def _process_runs(self, runs) -> str:
        """Process text runs with formatting"""
        html_parts = []

        for run in runs:
            text = self._escape_html(run.text) if run.text else ''
            if not text:
                continue

            # Apply formatting
            try:
                if run.bold:
                    text = f'<strong>{text}</strong>'
            except:
                pass

            try:
                if run.italic:
                    text = f'<em>{text}</em>'
            except:
                pass

            try:
                if run.underline:
                    text = f'<u>{text}</u>'
            except:
                pass

            # Check for font color
            try:
                if run.font and run.font.color and run.font.color.rgb:
                    rgb = run.font.color.rgb
                    # Handle different RGB formats
                    if hasattr(rgb, 'red'):
                        color = f'#{rgb.red:02x}{rgb.green:02x}{rgb.blue:02x}'
                    elif isinstance(rgb, (list, tuple)) and len(rgb) >= 3:
                        color = f'#{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}'
                    else:
                        color = None

                    if color:
                        text = f'<span style="color: {color};">{text}</span>'
            except Exception as e:
                logger.debug(f"Could not process font color: {e}")

            # Check for font size
            try:
                if run.font and run.font.size and run.font.size.pt:
                    size_pt = run.font.size.pt
                    text = f'<span style="font-size: {size_pt}pt;">{text}</span>'
            except:
                pass

            html_parts.append(text)

        return ''.join(html_parts)

    def convert_pdf_to_html(self, file_path: str) -> str:
        """Convert PDF to HTML with multiple strategies"""
        try:
            # Strategy 1: Try PyMuPDF for text-based PDFs
            html_content = self._convert_pdf_with_pymupdf(file_path)

            # Check if we got meaningful content
            if html_content and len(html_content.strip()) > 100:
                return html_content

            # Strategy 2: If no text or minimal text, use OCR
            if self.enable_ocr:
                logger.info("Attempting OCR conversion for PDF")
                return self._convert_pdf_with_ocr(file_path)

            return html_content

        except Exception as e:
            logger.error(f"Error converting PDF: {e}")
            return f"<p>Error converting PDF: {e}</p>"

    def _convert_pdf_with_pymupdf(self, file_path: str) -> str:
        """Convert PDF using PyMuPDF with formatting preservation"""
        html_parts = []

        try:
            pdf_document = fitz.open(file_path)

            for page_num, page in enumerate(pdf_document):
                # Extract text with formatting
                html_parts.append(f'<div class="pdf-page" data-page="{page_num + 1}">')

                # Get page text as HTML
                page_html = page.get_text("html")

                # Clean and enhance the HTML
                soup = BeautifulSoup(page_html, 'html.parser')

                # Remove unnecessary styles that might break layout
                for tag in soup.find_all(style=True):
                    style = tag.get('style', '')
                    # Keep only essential styles
                    cleaned_style = self._clean_pdf_styles(style)
                    if cleaned_style:
                        tag['style'] = cleaned_style
                    else:
                        del tag['style']

                # Extract images
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    try:
                        # Extract image
                        xref = img[0]
                        pix = fitz.Pixmap(pdf_document, xref)

                        if pix.n - pix.alpha < 4:  # GRAY or RGB
                            img_data = pix.tobytes("png")
                            img_b64 = base64.b64encode(img_data).decode()

                            # Insert image into HTML
                            img_tag = soup.new_tag('img')
                            img_tag['src'] = f"data:image/png;base64,{img_b64}"
                            img_tag['alt'] = f"Page {page_num + 1} Image {img_index + 1}"
                            img_tag['class'] = "pdf-image"

                            soup.append(img_tag)

                        pix = None  # Free resources

                    except Exception as e:
                        logger.warning(f"Failed to extract image: {e}")

                html_parts.append(str(soup))
                html_parts.append('</div>')

                # Add page break except for last page
                if page_num < len(pdf_document) - 1:
                    html_parts.append('<div class="page-break"></div>')

            pdf_document.close()

        except Exception as e:
            logger.error(f"PyMuPDF conversion error: {e}")
            return ""

        return '\n'.join(html_parts)

    def _clean_pdf_styles(self, style: str) -> str:
        """Clean PDF styles to keep only essential ones"""
        essential_properties = [
            'font-weight', 'font-style', 'text-decoration',
            'color', 'background-color', 'font-size',
            'text-align', 'margin', 'padding'
        ]

        try:
            style_dict = {}
            for prop in style.split(';'):
                if ':' in prop:
                    key, value = prop.split(':', 1)
                    key = key.strip()
                    if any(essential in key for essential in essential_properties):
                        style_dict[key] = value.strip()

            return '; '.join(f"{k}: {v}" for k, v in style_dict.items())
        except:
            return ""

    def _convert_pdf_with_ocr(self, file_path: str) -> str:
        """Convert PDF using OCR for scanned documents"""
        html_parts = []

        try:
            # Convert PDF pages to images
            images = convert_from_path(file_path, dpi=self.pdf_dpi)

            for i, image in enumerate(images):
                # Perform OCR
                text = pytesseract.image_to_string(image, config='--psm 6')

                # Also get layout information
                data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)

                # Build HTML with layout preservation
                page_html = self._build_html_from_ocr(data, image, i + 1)
                html_parts.append(page_html)

                # Add page break except for last page
                if i < len(images) - 1:
                    html_parts.append('<div class="page-break"></div>')

        except Exception as e:
            logger.error(f"OCR conversion error: {e}")
            return "<p>OCR conversion failed</p>"

        return '\n'.join(html_parts)

    def _build_html_from_ocr(self, ocr_data: Dict, image: Image, page_num: int) -> str:
        """Build HTML from OCR data with layout preservation"""
        html_parts = [f'<div class="ocr-page" data-page="{page_num}">']

        # Save the page image as base64
        img_buffer = io.BytesIO()
        image.save(img_buffer, format='PNG', quality=self.image_quality)
        img_b64 = base64.b64encode(img_buffer.getvalue()).decode()

        # Add background image for reference
        html_parts.append(f'''
        <div style="position: relative; width: 100%; margin-bottom: 20px;">
            <img src="data:image/png;base64,{img_b64}" alt="Page {page_num}" style="width: 100%; opacity: 0.1; position: absolute; top: 0; left: 0; z-index: -1;">
        ''')

        # Process OCR text with positioning
        current_block = []
        last_block_num = -1

        for i in range(len(ocr_data['text'])):
            if int(ocr_data['conf'][i]) > 0:  # Confidence threshold
                text = ocr_data['text'][i].strip()
                if text:
                    block_num = ocr_data['block_num'][i]

                    # New block
                    if block_num != last_block_num:
                        if current_block:
                            html_parts.append('<p>' + ' '.join(current_block) + '</p>')
                        current_block = [text]
                        last_block_num = block_num
                    else:
                        current_block.append(text)

        # Add last block
        if current_block:
            html_parts.append('<p>' + ' '.join(current_block) + '</p>')

        html_parts.append('</div></div>')

        return '\n'.join(html_parts)

    def convert_doc_to_html(self, file_path: str) -> str:
        """Convert DOC to HTML by first converting to DOCX, then to HTML"""
        temp_docx = self._convert_doc_to_docx(file_path)
        if temp_docx:
            try:
                html_content = self.convert_docx_to_html(temp_docx)
                os.unlink(temp_docx)
                return html_content
            except Exception as e:
                logger.error(f"DOCX to HTML conversion failed: {e}")
        return "<p>Unable to convert DOC file. Please try converting to DOCX first.</p>"

    def _check_libreoffice(self) -> bool:
        """Check if LibreOffice is available"""
        try:
            subprocess.run(['soffice', '--version'], capture_output=True, check=True)
            return True
        except:
            return False

    def _convert_doc_with_libreoffice(self, file_path: str) -> Optional[str]:
        """Convert DOC to HTML using LibreOffice"""
        try:
            temp_dir = tempfile.mkdtemp()
            output_file = os.path.join(temp_dir, 'output.html')

            cmd = [
                'soffice',
                '--headless',
                '--convert-to', 'html',
                '--outdir', temp_dir,
                file_path
            ]

            result = subprocess.run(cmd, capture_output=True, text=True)

            if result.returncode == 0 and os.path.exists(output_file):
                with open(output_file, 'r', encoding='utf-8') as f:
                    html_content = f.read()

                # Clean up
                shutil.rmtree(temp_dir)

                # Clean the HTML
                return self._clean_libreoffice_html(html_content)

        except Exception as e:
            logger.error(f"LibreOffice conversion failed: {e}")

        return None

    def _clean_libreoffice_html(self, html: str) -> str:
        """Clean HTML generated by LibreOffice"""
        soup = BeautifulSoup(html, 'html.parser')

        # Extract body content
        body = soup.find('body')
        if body:
            return str(body)

        return html

    def _convert_doc_with_word(self, file_path: str) -> Optional[str]:
        """Convert DOC to HTML using MS Word (Windows only)"""
        try:
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch('Word.Application')
            word.Visible = False

            doc = word.Documents.Open(os.path.abspath(file_path))

            # Save as filtered HTML
            temp_html = tempfile.mktemp(suffix='.html')
            doc.SaveAs(temp_html, FileFormat=10)  # 10 = Filtered HTML
            doc.Close()
            word.Quit()

            # Read the HTML
            with open(temp_html, 'r', encoding='utf-8') as f:
                html_content = f.read()

            # Clean up
            os.unlink(temp_html)

            # Clean the HTML
            return self._clean_word_html(html_content)

        except Exception as e:
            logger.error(f"Word conversion failed: {e}")

        return None

    def _clean_word_html(self, html: str) -> str:
        """Clean HTML generated by MS Word"""
        soup = BeautifulSoup(html, 'html.parser')

        # Remove Word-specific tags and attributes
        for tag in soup.find_all():
            # Remove Word-specific attributes
            for attr in list(tag.attrs.keys()):
                if attr.startswith('o:') or attr.startswith('v:'):
                    del tag[attr]

        # Remove empty paragraphs
        for p in soup.find_all('p'):
            if not p.text.strip():
                p.decompose()

        # Extract body content
        body = soup.find('body')
        if body:
            return str(body)

        return str(soup)

    def _convert_doc_to_docx(self, file_path: str) -> Optional[str]:
        """Convert DOC to DOCX using Word COM if available, else fallback to LibreOffice."""
        # Try Microsoft Word COM automation first
        if WINDOWS_AVAILABLE:
            try:
                pythoncom.CoInitialize()
                word = win32com.client.Dispatch('Word.Application')
                word.Visible = False

                doc = word.Documents.Open(os.path.abspath(file_path))
                temp_docx = tempfile.mktemp(suffix='.docx')
                doc.SaveAs(temp_docx, FileFormat=16)  # 16 = DOCX
                doc.Close()
                word.Quit()

                return temp_docx

            except Exception as e:
                logger.error(f"DOC to DOCX conversion failed (Word): {e}")

        # Fallback: Try LibreOffice
        try:
            import shutil
            import subprocess
            import tempfile
            temp_dir = tempfile.mkdtemp()
            output_file = os.path.join(temp_dir, os.path.splitext(os.path.basename(file_path))[0] + '.docx')
            cmd = [
                'soffice.com',
                '--headless',
                '--convert-to', 'docx',
                '--outdir', temp_dir,
                file_path
            ]
            result = subprocess.run(cmd, capture_output=True, text=True)
            if result.returncode == 0 and os.path.exists(output_file):
                return output_file
            else:
                logger.error(f"LibreOffice DOC to DOCX conversion failed: {result.stderr}")
        except Exception as e:
            logger.error(f"DOC to DOCX conversion failed (LibreOffice): {e}")
        return None

    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters"""
        if not text:
            return ""
        return (text
                .replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;')
                .replace("'", '&#39;'))

    def process_document(self, file_path: str) -> Dict[str, Any]:
        """Process a single document and convert to HTML"""
        result = {
            'filename': os.path.basename(file_path),
            'input_path': os.path.abspath(file_path),
            'timestamp': datetime.now().isoformat(),
            'status': 'processing'
        }

        try:
            # Detect file type
            ext = Path(file_path).suffix.lower()

            # Convert based on file type
            if ext == '.pdf':
                html_content = self.convert_pdf_to_html(file_path)
            elif ext == '.docx':
                html_content = self.convert_docx_to_html(file_path)
            elif ext == '.doc':
                html_content = self.convert_doc_to_html(file_path)
            else:
                raise ValueError(f"Unsupported file type: {ext}")

            # Create full HTML document
            full_html = self.html_template.format(
                title=Path(file_path).stem,
                styles=self.base_css,
                content=html_content,
                scripts=self.base_js
            )

            # Save HTML file
            output_filename = Path(file_path).stem + '.html'
            output_path = os.path.join(self.output_dir, output_filename)

            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_html)

            # Create metadata
            result.update({
                'status': 'success',
                'output_path': output_path,
                'html_size': len(full_html),
                'file_size': os.path.getsize(file_path)
            })

            # Save metadata
            metadata_path = os.path.join(self.output_dir, Path(file_path).stem + '_metadata.json')
            with open(metadata_path, 'w', encoding='utf-8') as f:
                json.dump(result, f, indent=2)

        except Exception as e:
            logger.error(f"Error processing {file_path}: {e}")
            result.update({
                'status': 'failed',
                'error': str(e)
            })

        return result

    def batch_process(self, input_folder: str, max_workers: int = 4) -> Dict[str, Any]:
        """Process multiple documents in parallel"""
        from concurrent.futures import ThreadPoolExecutor, as_completed

        # Find all documents
        documents = []
        for ext in ['.pdf', '.doc', '.docx']:
            documents.extend(Path(input_folder).glob(f'*{ext}'))
            documents.extend(Path(input_folder).glob(f'*{ext.upper()}'))

        documents = list(set(documents))  # Remove duplicates
        logger.info(f"Found {len(documents)} documents to process")

        results = {
            'total': len(documents),
            'successful': 0,
            'failed': 0,
            'documents': []
        }

        # Process documents in parallel
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_to_doc = {executor.submit(self.process_document, str(doc)): doc
                             for doc in documents}

            for future in as_completed(future_to_doc):
                doc = future_to_doc[future]
                try:
                    result = future.result()
                    results['documents'].append(result)

                    if result['status'] == 'success':
                        results['successful'] += 1
                        print(f"✓ Converted: {doc.name}")
                    else:
                        results['failed'] += 1
                        print(f"✗ Failed: {doc.name} - {result.get('error', 'Unknown error')}")

                except Exception as e:
                    logger.error(f"Unexpected error processing {doc}: {e}")
                    results['failed'] += 1
                    results['documents'].append({
                        'filename': doc.name,
                        'status': 'failed',
                        'error': str(e)
                    })

        # Save batch summary
        summary_path = os.path.join(self.output_dir, 'batch_summary.json')
        with open(summary_path, 'w', encoding='utf-8') as f:
            json.dump(results, f, indent=2)

        return results


class TextExtractor:
    """Extract plain text from PDF, DOCX, and DOC files."""
    def __init__(self, antiword_path: str = r'C:\antiword\antiword.exe'):
        self.antiword_path = antiword_path

    def extract_text_from_pdf(self, file_path: str) -> str:
        import pdfplumber
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            text = f"Error extracting PDF text: {e}"
        return text

    def extract_text_from_docx(self, file_path: str) -> str:
        from docx import Document
        text = ""
        try:
            doc = Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except Exception as e:
            text = f"Error extracting DOCX text: {e}"
        return text

    def extract_text_from_doc(self, file_path: str) -> str:
        import os
        import subprocess
        text = ""
        # Try Windows COM automation first
        try:
            import win32com.client
            import pythoncom
            pythoncom.CoInitialize()
            import tempfile
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_docx = temp_file.name
            word = win32com.client.Dispatch('Word.Application')
            word.Visible = False
            doc = word.Documents.Open(os.path.abspath(file_path))
            doc.SaveAs(os.path.abspath(temp_docx), 16)  # 16 = docx
            doc.Close()
            word.Quit()
            text = self.extract_text_from_docx(temp_docx)
            os.remove(temp_docx)
            if text.strip():
                return text
        except Exception:
            pass
        # Fallback to antiword
        try:
            if os.path.exists(self.antiword_path):
                result = subprocess.run(
                    [self.antiword_path, file_path],
                    capture_output=True,
                    text=True,
                    encoding='utf-8',
                    errors='replace'
                )
                if result.returncode == 0:
                    return result.stdout
        except Exception:
            pass
        # Fallback to mammoth
        try:
            import mammoth
            with open(file_path, 'rb') as doc_file:
                result = mammoth.extract_raw_text(doc_file)
                if result.value:
                    return result.value
        except Exception:
            pass
        return "Error: Could not extract text from DOC file"

    def extract_text(self, file_path: str) -> str:
        ext = str(file_path).lower()
        if ext.endswith('.pdf'):
            return self.extract_text_from_pdf(file_path)
        elif ext.endswith('.docx'):
            return self.extract_text_from_docx(file_path)
        elif ext.endswith('.doc'):
            return self.extract_text_from_doc(file_path)
        else:
            return "Unsupported file type for text extraction."


class EnhancedHTMLGenerator:
    """Generate enhanced HTML with advanced features"""

    @staticmethod
    def create_navigation_menu(sections: List[str]) -> str:
        """Create a navigation menu for the resume"""
        nav_html = '<nav class="resume-nav">\n<ul>\n'
        for section in sections:
            section_id = section.lower().replace(' ', '-')
            nav_html += f'<li><a href="#{section_id}">{section}</a></li>\n'
        nav_html += '</ul>\n</nav>'
        return nav_html

    @staticmethod
    def add_download_buttons() -> str:
        """Add download buttons for different formats"""
        return """
        <div class="download-buttons">
            <button onclick="printResume()" class="btn btn-print">Print</button>
            <button onclick="downloadAsPDF()" class="btn btn-pdf">Download PDF</button>
            <button onclick="downloadAsWord()" class="btn btn-word">Download Word</button>
        </div>
        """

    @staticmethod
    def enhance_contact_section(html: str) -> str:
        """Enhance contact information with icons and links"""
        # This would parse contact info and add appropriate icons
        # For now, returning as-is
        return html

    @staticmethod
    def add_responsive_styles() -> str:
        """Additional responsive styles for modern devices"""
        return """
        /* Modern responsive design */
        @media (prefers-color-scheme: dark) {
            body { background-color: #1a1a1a; color: #e0e0e0; }
            .resume-container { background-color: #2a2a2a; }
            h1, h2, h3, h4, h5, h6 { color: #4a9eff; }
        }

        /* Mobile optimizations */
        @media screen and (max-width: 480px) {
            .resume-container { padding: 15px; }
            h1 { font-size: 1.8em; }
            h2 { font-size: 1.4em; }
            table { font-size: 0.9em; }
        }

        /* Print optimizations */
        @media print {
            .download-buttons, .resume-nav { display: none; }
            body { font-size: 11pt; }
            .page-break { page-break-before: always; }
        }

        /* Interactive elements */
        .btn {
            padding: 10px 20px;
            margin: 5px;
            border: none;
            border-radius: 5px;
            cursor: pointer;
            font-size: 14px;
            transition: all 0.3s ease;
        }

        .btn-print {
            background-color: #3498db;
            color: white;
        }

        .btn-print:hover {
            background-color: #2980b9;
        }

        .btn-pdf {
            background-color: #e74c3c;
            color: white;
        }

        .btn-pdf:hover {
            background-color: #c0392b;
        }

        .btn-word {
            background-color: #2ecc71;
            color: white;
        }

        .btn-word:hover {
            background-color: #27ae60;
        }
        """


def main():
    """Command-line interface for the converter"""
    import argparse

    parser = argparse.ArgumentParser(description='Convert resumes to high-quality HTML')
    parser.add_argument('input', help='Input file or folder path')
    parser.add_argument('-o', '--output', default='html_output', help='Output directory')
    parser.add_argument('--no-ocr', action='store_true', help='Disable OCR for PDFs')
    parser.add_argument('--quality', type=int, default=85, help='Image quality (1-100)')
    parser.add_argument('--dpi', type=int, default=150, help='PDF conversion DPI')
    parser.add_argument('--workers', type=int, default=4, help='Number of parallel workers')
    parser.add_argument('--batch', action='store_true', help='Process entire folder')

    args = parser.parse_args()

    # Initialize converter
    converter = HighQualityHTMLConverter(
        output_dir=args.output,
        enable_ocr=not args.no_ocr,
        image_quality=args.quality,
        pdf_dpi=args.dpi
    )

    if args.batch or os.path.isdir(args.input):
        # Batch processing
        print(f"Starting batch conversion of {args.input}")
        results = converter.batch_process(args.input, max_workers=args.workers)

        print(f"\nConversion complete!")
        print(f"Total files: {results['total']}")
        print(f"Successful: {results['successful']}")
        print(f"Failed: {results['failed']}")
        print(f"Output directory: {args.output}")

    else:
        # Single file processing
        if not os.path.isfile(args.input):
            print(f"Error: File not found: {args.input}")
            sys.exit(1)

        print(f"Converting {args.input}")
        result = converter.process_document(args.input)

        if result['status'] == 'success':
            print(f"✓ Successfully converted to: {result['output_path']}")
        else:
            print(f"✗ Conversion failed: {result.get('error', 'Unknown error')}")


if __name__ == "__main__":
    main()